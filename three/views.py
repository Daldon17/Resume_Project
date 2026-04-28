import re
import PyPDF2
import docx
from django.shortcuts import render
from .forms import ResumeForm
from .models import Resume
from django.http import HttpResponse
from django.template.loader import get_template
from xhtml2pdf import pisa
import pdfplumber

# EXPERIENCE EXTRACTION

def extract_required_experience(job_text):
    job_text = job_text.lower()
    match = re.search(r'(\d+)\s*(?:to|-)\s*(\d+)\s*(years|yrs)', job_text)
    if match:
        return int(match.group(1)), int(match.group(2))
    match = re.search(r'(?:minimum|min)\s*(\d+)\s*(years|yrs)', job_text)
    if match:
        return int(match.group(1)), int(match.group(1))
    return None, None

def extract_resume_experience(text):
    text = text.lower()
    if "fresher" in text or "entry level" in text:
        return 0
    match = re.search(r'(\d+)\+?\s*(years|yrs)', text)
    if match:
        return int(match.group(1))
    if "intern" in text:
        return 0.5
    return 0

# JOB DATA (ATS KNOWLEDGE BASE)

JOB_DATA = {
    'python_dev': {
        'keywords': ['python', 'django', 'api', 'sql', 'flask'],
        'job_titles': ['python developer', 'backend developer'],
        'education': ['bsc', 'b.tech', 'computer science'],
        'certifications': ['aws', 'azure', 'google cloud', 'python certification']
    },
    'frontend_dev': {
        'keywords': ['html', 'css', 'javascript', 'react', 'bootstrap'],
        'job_titles': ['frontend developer', 'web developer'],
        'education': ['bsc', 'bca', 'be'],
        'certifications': ['figma', 'ui ux', 'frontend certification']
    },
    'data_analyst': {
        'keywords': ['python', 'sql', 'excel', 'powerbi', 'tableau'],
        'job_titles': ['data analyst', 'business analyst'],
        'education': ['statistics', 'mathematics', 'bsc'],
        'certifications': [
            'powerbi certification',
            'tableau certification',
            'google data analytics'
        ]
    }
}

SOFT_SKILLS = {
    'python_dev': ['problem solving', 'analytical thinking', 'attention to detail'],
    'frontend_dev': ['creativity', 'communication', 'collaboration'],
    'data_analyst': ['analytical thinking', 'critical thinking', 'communication']
}

ROLE_BASED_PROJECT_KEYWORDS = {
    'python_dev': ['api', 'django', 'flask', 'backend', 'rest', 'database', 'automation', 'scraping'],
    'frontend_dev': ['ui', 'ux', 'frontend', 'react', 'html', 'css', 'javascript', 'bootstrap', 'website'],
    'data_analyst': ['analysis', 'dashboard', 'powerbi', 'tableau', 'excel', 'data', 'visualization', 'report']
}

ROLE_BASED_INTERNSHIP_KEYWORDS = {
    'python_dev': ['backend intern', 'python intern', 'django intern', 'software intern'],
    'frontend_dev': ['ui intern', 'frontend intern', 'web intern', 'design intern'],
    'data_analyst': ['data intern', 'analyst intern', 'business analyst intern']
}

# ROLE DETECTION

def detect_role(job_text):
    job_text = job_text.lower()
    role_keywords = {
        'python_dev': ['python', 'django', 'flask', 'api', 'backend'],
        'frontend_dev': ['html', 'css', 'javascript', 'react', 'frontend'],
        'data_analyst': ['sql', 'excel', 'powerbi', 'tableau', 'data']
    }
    scores = {role: sum(1 for k in keywords if k in job_text)
              for role, keywords in role_keywords.items()}
    if max(scores.values()) == 0:
        return 'python_dev'
    return max(scores, key=scores.get)

# RESUME TEXT EXTRACTION

def extract_text(file):
    text = ""
    try:
        if file.name.endswith('.pdf'):
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                content = page.extract_text()
                if content:
                    text += content
        elif file.name.endswith('.docx'):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text
    except Exception as e:
        print("File Read Error:", e)
    return text.lower().replace("-", " ")

# KEYWORD MATCHING

def keyword_match(text, keywords):
    matched = []
    missing = []
    for keyword in keywords:
        pattern = rf'\b{re.escape(keyword.lower())}\b'
        if re.search(pattern, text):
            matched.append(keyword)
        else:
            missing.append(keyword)
    return matched, missing

# ATS SCORING ENGINE

def calculate_ats_score(text, role, job_description):
    if not text.strip():
        return 0, {}
    data = JOB_DATA.get(role, {})
    keywords = data.get('keywords', [])
    job_titles = data.get('job_titles', [])
    education = data.get('education', [])
    certifications = data.get('certifications', [])

    # KEYWORDS

    matched, missing = keyword_match(text, keywords)
    keyword_score = int((len(matched) / len(keywords)) * 30) if keywords else 0
    keyword_percent = (keyword_score / 30) * 100

    # JOB TITLE 

    matched_titles = []
    for t in job_titles:
        words = t.split()
        if all(w in text for w in words):
            matched_titles = [t]
            break
    if matched_titles:
        job_title_score = 10
        job_title_missing = []
    else:
        job_title_score = 0
        job_title_missing = job_titles

    # SOFT SKILLS

    soft_skills = SOFT_SKILLS.get(role, [])
    soft_matched = [s for s in soft_skills if re.search(rf'\b{s}\b', text)]
    soft_missing = list(set(soft_skills) - set(soft_matched))

    skills_score = int((len(soft_matched) / len(soft_skills)) * 10) if soft_skills else 0

    # PROJECTS
   
    project_keywords = ROLE_BASED_PROJECT_KEYWORDS.get(role, [])
    project_matches = [p for p in project_keywords if p in text]
    if project_matches:
        project_score = min(10, len(project_matches) * 3 + 4)
        project_suggestion = ""
    else:
        project_score = 0
        project_suggestion = "Add relevant projects"

    # INTERNSHIPS
   
    internship_keywords = ROLE_BASED_INTERNSHIP_KEYWORDS.get(role, [])
    internship_found = bool(re.search(r'intern|internship|trainee', text))
    intern_matches = [i for i in internship_keywords if i in text]
    if intern_matches:
        internship_score = min(10, len(intern_matches) * 4 + 3)
        internship_suggestion = ""
    elif internship_found:
        internship_score = 5
        internship_suggestion = "Add role-specific internship experience"
    else:
        internship_score = 0
        internship_suggestion = "Add internship experience"

    # EXPERIENCE 
  
    req_min, req_max = extract_required_experience(job_description)
    resume_exp = extract_resume_experience(text)
    exp_matched = []
    exp_missing = []
    exp_suggestion = ""
    if req_min is not None:
        if resume_exp == 0:
            exp_score = 6
            exp_matched = ["Fresher"]
            exp_suggestion = "Relevant experience would be considered a strong advantage."
        elif req_max is not None and req_min <= resume_exp <= req_max:
            exp_score = 10
            exp_matched = [f"{resume_exp} years matches requirement"]
        elif resume_exp < req_min:
            exp_score = 3
            exp_missing = [f"Required {req_min}-{req_max} years"]
            exp_suggestion = f"Need at least {req_min} years experience"
        else:
            exp_score = 8
            exp_matched = [f"{resume_exp} years (overqualified)"]
    else:
        if resume_exp > 0:
            exp_score = 8
            exp_matched = [f"{resume_exp} years"]
        else:
            exp_score = 5
            exp_missing = ["Experience not clearly mentioned"]

    # EDUCATION

    matched_edu = [e for e in education if e in text]
    if matched_edu:
        edu_score = 10
        edu_missing = []
        edu_suggestion = ""
    else:
        edu_score = 0
        edu_missing = education
        edu_suggestion = "Add education"

    # CERTIFICATIONS

    cert_text_found = bool(re.search(r'certification|certified|course|training', text))
    cert_matched = [c for c in certifications if c in text]
    cert_missing = list(set(certifications) - set(cert_matched))
    if cert_matched:
        cert_score = int((len(cert_matched) / len(certifications)) * 10)
    elif cert_text_found:
        
        cert_score = 5
        cert_missing = []
    else:
        cert_score = 0
    cert_suggestion = "Add certifications" if cert_score == 0 else ""

    # TOTAL SCORE

    total_score = min((
        keyword_score +          
        job_title_score +        
        skills_score +           
        project_score +          
        internship_score +      
        exp_score +              
        edu_score +              
        cert_score               
    ), 100)

    # BREAKDOWN

    breakdown = {
'keywords': {
    'score': keyword_score,
    'percent': keyword_percent,
    'matched': matched,
    'missing': missing,
    'suggestion': f"Add: {', '.join(missing)}" if missing else ""
},

        'soft_skills': {
            'score': skills_score,
            'percent': (skills_score / 10) * 100,
            'matched': soft_matched,
            'missing': soft_missing,
            'suggestion': "Improve soft skills" if skills_score < 10 else ""
        },

        'projects': {
            'score': project_score,
            'percent': (project_score / 10) * 100,
            'matched': project_matches,
            'missing': [] if project_matches else project_keywords,
            'suggestion': project_suggestion
        },

        'internships': {
            'score': internship_score,
            'percent': (internship_score / 10) * 100,
            'matched': intern_matches,
            'missing': [] if intern_matches else internship_keywords,
            'suggestion': internship_suggestion
        },

        'experience': {
            'score': exp_score,
            'percent': (exp_score / 10) * 100,
            'matched': exp_matched,
            'missing': exp_missing,
            'suggestion': exp_suggestion
        },

        'education': {
            'score': edu_score,
            'percent': (edu_score / 10) * 100,
            'matched': matched_edu,
            'missing': edu_missing,
            'suggestion': edu_suggestion
        },

        'certifications': {
            'score': cert_score,
            'percent': (cert_score / 10) * 100,
            'matched': cert_matched,
            'missing': cert_missing,
            'suggestion': cert_suggestion
        },

    'job_title': {
        'score': job_title_score,

        # ✅ 100% if matched, else 0%
        'percent': 100 if job_title_score == 10 else 0,

        'matched': matched_titles,
        'missing': job_title_missing,

        'suggestion': (
            "" if job_title_score == 10
            else "Add job title like: " + ", ".join(job_titles)
        )
    }
    }

    return total_score, breakdown

# MAIN VIEW

def upload_resume(request):
    result = None
    if request.method == 'POST':
        form = ResumeForm(request.POST, request.FILES)
        if form.is_valid():
            resume = form.save(commit=False)
            file = form.cleaned_data['file']
            text = extract_text(file)
            job_description = request.POST.get('job_description', '')
            role = detect_role(job_description)
            score, breakdown = calculate_ats_score(text, role, job_description)
            for section in breakdown.values():
                section['missing_percent'] = 100 - section.get('percent', 0)
            resume.score = score
            resume.matched_keywords = ", ".join(breakdown['keywords']['matched'])
            resume.missing_keywords = ", ".join(breakdown['keywords']['missing'])
            resume.skills_score = breakdown['soft_skills']['score']
            resume.experience_score = breakdown['experience']['score']
            resume.education_score = breakdown['education']['score']
            resume.certifications_score = breakdown['certifications']['score']
            resume.job_title_score = breakdown['job_title']['score']
            resume.save()
            result = {
                'score': score,
                'breakdown': breakdown,
                'detected_role': role
            }
    else:
        form = ResumeForm()
    return render(request, 'upload.html', {'form': form, 'result': result})

def select_template(request):
    templates = [
        {'name': 'Template 1', 'file': 'template1.html', 'image': 'images/t1.png'},
        {'name': 'Template 2', 'file': 'template2.html', 'image': 'images/t2.png'},
        {'name': 'Template 3', 'file': 'template3.html', 'image': 'images/t3.png'},
        {'name': 'Template 4', 'file': 'template4.html', 'image': 'images/t4.png'},
    ]
    return render(request, 'select_template.html', {'templates': templates})

def create_resume(request, template_name):
    return render(request, 'resume_form.html', {'template': template_name})

def preview_resume(request):
    if request.method == 'POST':
        data = request.POST
        template = request.POST.get('template')

        return render(request, f'resumes/{template}', {'data': data})

def download_resume(request):
    if request.method == 'POST':
        template = request.POST.get('template')
        data = request.POST

        template_path = f'resumes/{template}'
        template = get_template(template_path)
        html = template.render({
    'data': data,
    'is_pdf': True   
})
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="resume.pdf"'

        pisa.CreatePDF(html, dest=response)
        return response


# ===== NEW API CODE START =====
from pathlib import Path

from django.http import JsonResponse

try:
    import pdfplumber
except ImportError:  
    pdfplumber = None

try:
    from rest_framework import status
    from rest_framework.decorators import api_view, parser_classes
    from rest_framework.parsers import FormParser, MultiPartParser
    from rest_framework.response import Response
except ImportError:  
    class _StatusFallback:
        HTTP_200_OK = 200
        HTTP_201_CREATED = 201
        HTTP_400_BAD_REQUEST = 400
        HTTP_404_NOT_FOUND = 404
        HTTP_500_INTERNAL_SERVER_ERROR = 500

    status = _StatusFallback()

    def api_view(_methods):
        def decorator(func):
            return func
        return decorator

    def parser_classes(_parsers):
        def decorator(func):
            return func
        return decorator

    MultiPartParser = FormParser = object
    Response = JsonResponse

from .serializers import ResumeSerializer

_legacy_extract_text = extract_text
RESUME_TEMPLATE_DIR = Path(__file__).resolve().parent / "templates" / "resumes"
API_ENDPOINTS = {
    "root": {"method": "GET", "url": "/api/"},
    "health": {"method": "GET", "url": "/api/health/"},
    "upload_resume": {"method": "POST", "url": "/api/upload-resume/"},
    "analyze_resume": {"method": "POST", "url": "/api/analyze/"},
    "resume_results": {"method": "GET", "url": "/api/results/<id>/"},
    "resume_templates": {"method": "GET", "url": "/api/templates/"},
}
COMMON_JOB_STOPWORDS = {
    "a", "an", "and", "are", "as", "at", "be", "build", "by", "for", "from",
    "has", "have", "in", "into", "is", "it", "knowledge", "looking", "must",
    "of", "on", "or", "our", "role", "should", "strong", "team", "that",
    "the", "their", "this", "to", "using", "we", "will", "with", "you",
    "your",
}


def _normalize_api_text(text):
    return re.sub(r"\s+", " ", (text or "").lower().replace("-", " ")).strip()


def _match_keyword_in_text(keyword, text):
    pattern = rf"(?<!\w){re.escape(keyword.lower())}(?!\w)"
    return bool(re.search(pattern, text))


def _extract_keywords_from_job_description(job_description):
    raw_keywords = re.findall(r"[A-Za-z][A-Za-z0-9.+#/-]*", job_description.lower())
    extracted_keywords = []
    seen_keywords = set()

    for keyword in raw_keywords:
        cleaned_keyword = keyword.strip(".-/")
        if (
            len(cleaned_keyword) < 2
            or cleaned_keyword.isdigit()
            or cleaned_keyword in COMMON_JOB_STOPWORDS
            or cleaned_keyword in seen_keywords
        ):
            continue

        extracted_keywords.append(cleaned_keyword)
        seen_keywords.add(cleaned_keyword)

    return extracted_keywords


def _template_display_name(template_stem):
    match = re.search(r"(\d+)$", template_stem)
    if match:
        return f"Template {match.group(1)}"
    return template_stem.replace("_", " ").title()


def _get_error_message(errors, default_message="Invalid request data."):
    if isinstance(errors, dict):
        for value in errors.values():
            if isinstance(value, (list, tuple)) and value:
                return str(value[0])
            if value:
                return str(value)
    elif isinstance(errors, (list, tuple)) and errors:
        return str(errors[0])
    elif errors:
        return str(errors)

    return default_message


def extract_text(file_path):
    if not file_path:
        return ""

    file_name = getattr(file_path, "name", str(file_path))
    file_extension = Path(file_name).suffix.lower()
    is_file_object = hasattr(file_path, "read")

    if is_file_object:
        try:
            file_path.seek(0)
        except (AttributeError, OSError):
            pass

    try:
        if file_extension == ".pdf":
            if pdfplumber is not None:
                extracted_pages = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            extracted_pages.append(page_text)
                return _normalize_api_text(" ".join(extracted_pages))

            return _legacy_extract_text(file_path)

        if file_extension == ".docx":
            document = docx.Document(file_path)
            paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            return _normalize_api_text(" ".join(paragraphs))

        if is_file_object:
            return _legacy_extract_text(file_path)

        raise ValueError("Unsupported file format. Only PDF and DOCX files are supported.")
    except Exception as exc:
        if is_file_object:
            try:
                return _legacy_extract_text(file_path)
            except Exception as legacy_exc:
                raise ValueError(
                    f"Unable to extract text from the uploaded file: {legacy_exc}"
                ) from exc
        raise ValueError(f"Unable to extract text from the uploaded file: {exc}") from exc
    finally:
        if is_file_object:
            try:
                file_path.seek(0)
            except (AttributeError, OSError):
                pass


def analyze_resume(text, job_description):
    normalized_text = _normalize_api_text(text)
    normalized_job_description = _normalize_api_text(job_description)

    if not normalized_text:
        raise ValueError("No readable text could be extracted from the resume.")

    if not normalized_job_description:
        raise ValueError("job_description is required.")

    detected_role = detect_role(normalized_job_description)
    extracted_keywords = _extract_keywords_from_job_description(normalized_job_description)
    fallback_keywords = JOB_DATA.get(detected_role, {}).get("keywords", [])
    keywords = extracted_keywords or fallback_keywords

    matched_keywords = []
    missing_keywords = []
    for keyword in keywords:
        if _match_keyword_in_text(keyword, normalized_text):
            matched_keywords.append(keyword)
        else:
            missing_keywords.append(keyword)

    total_score, breakdown = calculate_ats_score(
        normalized_text,
        detected_role,
        normalized_job_description,
    )

    keyword_percent = round((len(matched_keywords) / len(keywords)) * 100, 2) if keywords else 0
    keyword_section_score = int(round((len(matched_keywords) / len(keywords)) * 30)) if keywords else 0

    analysis_breakdown = {
        "keywords": {
            "score": keyword_section_score,
            "percent": keyword_percent,
            "matched": matched_keywords,
            "missing": missing_keywords,
            "suggestion": f"Add: {', '.join(missing_keywords)}" if missing_keywords else "",
        },
        "skills": breakdown.get("soft_skills", {}),
        "experience": breakdown.get("experience", {}),
        "education": breakdown.get("education", {}),
        "certifications": breakdown.get("certifications", {}),
        "job_title": breakdown.get("job_title", {}),
        "projects": breakdown.get("projects", {}),
        "internships": breakdown.get("internships", {}),
    }

    return {
        "score": total_score,
        "detected_role": detected_role,
        "matched_keywords": matched_keywords,
        "missing_keywords": missing_keywords,
        "section_scores": {
            "skills": breakdown.get("soft_skills", {}).get("score", 0),
            "experience": breakdown.get("experience", {}).get("score", 0),
            "education": breakdown.get("education", {}).get("score", 0),
        },
        "breakdown": analysis_breakdown,
    }

@api_view(["GET"])
def api_root(request):
    return Response(
        {"endpoints": API_ENDPOINTS},
        status=status.HTTP_200_OK,
    )

@api_view(["GET"])
def api_health(request):
    return Response(
        {"status": "OK"},
        status=status.HTTP_200_OK,
    )

@api_view(["POST"])
@parser_classes([MultiPartParser, FormParser])
def api_upload_resume(request):
    print("api_upload_resume request.FILES:", request.FILES)
    print("api_upload_resume request.data:", request.data)
    try:
        uploaded_file = request.FILES.get("file")
        if not uploaded_file:
            return Response(
                {
                    "error": 'A PDF resume file is required. Send the file as form-data using the "file" key.'
                },
                status=status.HTTP_400_BAD_REQUEST,
            )
        serializer = ResumeSerializer(
            data={"file": uploaded_file},
            context={"request": request},
        )
        if not serializer.is_valid():
            return Response(
                {
                    "error": _get_error_message(serializer.errors),
                    "details": serializer.errors,
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        resume = serializer.save()
        return Response(
            {
                "id": resume.id,
                "message": "Resume uploaded successfully.",
                "file_name": uploaded_file.name,
            },
            status=status.HTTP_201_CREATED,
        )
    except Exception as exc:
        return Response(
            {
                "error": "An unexpected error occurred while uploading the resume.",
                "details": str(exc),
            },
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


@api_view(["POST"])
def api_analyze_resume(request):
    print("api_analyze_resume request.FILES:", request.FILES)
    print("api_analyze_resume request.data:", request.data)

    try:
        resume_id = request.data.get("resume_id")
        job_description = request.data.get("job_description", "")

        if not resume_id:
            return Response(
                {"error": "resume_id is required."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        if not job_description:
            return Response(
                {"error": "job_description is required."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            resume = Resume.objects.get(pk=resume_id)
        except Resume.DoesNotExist:
            return Response(
                {"error": "Resume not found."},
                status=status.HTTP_404_NOT_FOUND,
            )

        if not resume.file:
            return Response(
                {"error": "No resume file is attached to this record."},
                status=status.HTTP_404_NOT_FOUND,
            )

        try:
            file_path = resume.file.path
            if not Path(file_path).exists():
                raise FileNotFoundError("Resume file not found on disk.")
            resume_text = extract_text(file_path)
        except (ValueError, FileNotFoundError, NotImplementedError) as exc:
            return Response(
                {"error": str(exc)},
                status=status.HTTP_404_NOT_FOUND,
            )
        except Exception as exc:
            return Response(
                {"error": f"Unable to process the resume file: {exc}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        try:
            analysis = analyze_resume(resume_text, job_description)
        except ValueError as exc:
            return Response(
                {"error": str(exc)},
                status=status.HTTP_400_BAD_REQUEST,
            )

        resume.score = analysis["score"]
        resume.matched_keywords = ", ".join(analysis["matched_keywords"])
        resume.missing_keywords = ", ".join(analysis["missing_keywords"])
        resume.skills_score = analysis["section_scores"]["skills"]
        resume.experience_score = analysis["section_scores"]["experience"]
        resume.education_score = analysis["section_scores"]["education"]
        resume.certifications_score = analysis["breakdown"]["certifications"].get("score", 0)
        resume.job_title_score = analysis["breakdown"]["job_title"].get("score", 0)
        resume.save(
            update_fields=[
                "score",
                "matched_keywords",
                "missing_keywords",
                "skills_score",
                "experience_score",
                "education_score",
                "certifications_score",
                "job_title_score",
            ]
        )

        return Response(
            {
                "resume_id": resume.id,
                "analysis": analysis,
                "resume": ResumeSerializer(resume, context={"request": request}).data,
            },
            status=status.HTTP_200_OK,
        )
    except Exception as exc:
        return Response(
            {
                "error": "An unexpected error occurred while analyzing the resume.",
                "details": str(exc),
            },
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


@api_view(["GET"])
def api_resume_results(request, id):
    try:
        resume = Resume.objects.get(pk=id)
    except Resume.DoesNotExist:
        return Response(
            {"error": "Resume not found."},
            status=status.HTTP_404_NOT_FOUND,
        )

    serializer = ResumeSerializer(resume, context={"request": request})
    return Response(serializer.data, status=status.HTTP_200_OK)


@api_view(["GET"])
def api_resume_templates(request):
    templates = []
    if RESUME_TEMPLATE_DIR.exists():
        for template_path in sorted(RESUME_TEMPLATE_DIR.glob("*.html")):
            templates.append(
                {
                    "name": _template_display_name(template_path.stem),
                    "file": template_path.name,
                }
            )

    return Response(
        {"templates": templates},
        status=status.HTTP_200_OK,
    )
# ===== NEW API CODE END =====
